import streamlit as st
from docx import Document
from docx.shared import Cm

png_listesi = ["wordcloud.png", "ngram.png", "anahtar.png", "duygu.png", "ilk50.png", "benzerlik.png", "ağ.png", "merkezi.png", "hubs.png", "authority.png", "birliktelik.png", "takip.png", "üçkelime.png", "aradalılık.png", "kelimetürü.png", "isimler.png", "fiiller.png", "sıfatlar.png", "tekrareden.png"]

baslik_listesi = ["Kelime Bulutu", "İkili Ngram Kelime Bulutu", "Anahtar Kelime Bulutu", "Duygusal Analiz", "En Önemli İlk 50 Kelime İlişki Haritası", "Kelimelerin Yakınlıklarına Göre Sınıflanmış Haritası", "Kelime Ağ Analizi", "Merkezilik Ağ Analizi", "İlişki Haritası (Hubs)", "İlişki Haritası (Authority)", "Cümlede Birlikte Bulunma (İsim & Sıfat)", "Bir Diğerini Takip Eden Kelimeler (İsim & Sıfat)", "Üç Kelime İçinde Birlikte Bulunma (İsim & Sıfat)", "Güç, Aradalık ve Ağırlık Merkezi", "Kelime Tipleri", "İsimler", "Fiiller", "Sıfatlar", "Sık Tekrar Eden Kelime Grafiği"]

st.title('Grafikler ve Başlıklar')

doc = Document()

sayac = 0
for baslik, png in zip(baslik_listesi, png_listesi):
    st.header(baslik)
    st.image(png, caption=baslik, use_column_width=True)
    
    doc.add_heading(baslik, level=2)
    doc.add_picture(png, width=Cm(14), height=Cm(8.5))
    
    if sayac % 2 != 1:
        doc.add_paragraph().paragraph_format.space_after = 0
    sayac += 1
    
    if sayac % 2 == 0:
        doc.add_page_break()

st.write("Belgeler tarayıcıda gösterildi.")
